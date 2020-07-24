#!/usr/bin/env python3
# Takes a list of names and makes invites via Custom styles in .docx files 

import docx

with open('guests.txt') as f:
    names = f.readlines()
    document = docx.Document()

    for name in names:
        name = name.strip()

        # Invite contents
        document.add_paragraph('It would be a pleasure to have the company of', style='Custom 1')
        document.add_paragraph(name, style='Custom 2')
        document.add_paragraph('at 11010 Memory Lane on the Evening of', style='Custom 3')
        document.add_paragraph('April 1st', style='Custom 4')
        document.add_paragraph("at 7 o'clock", style='Custom 5')

        # New page for each guest
        document.add_page_break()

    document.save('invites.docx')

    print("Custom invites saved in 'invites.docx'")